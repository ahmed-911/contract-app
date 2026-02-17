import io
import os
import zipfile
import tempfile
import subprocess
import pandas as pd
import streamlit as st
from docx import Document

st.set_page_config(page_title="توليد عقود من Excel", layout="centered")
st.title("for you")
st.image("logo.png", width=200)


# ------------------ LibreOffice (Linux/Docker) ------------------
SOFFICE = os.environ.get("SOFFICE_PATH", "soffice")  # داخل Docker يكون "soffice"
PDF_AVAILABLE = True  # نفترض موجود لأن Dockerfile يثبته

# ---------- Word replacement helpers (حتى لو placeholder مقسوم runs) ----------
def _replace_in_paragraph(paragraph, mapping):
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = full_text
    for k, v in mapping.items():
        replaced = replaced.replace(k, v)

    if replaced != full_text:
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph.runs[i].text = ""
        if paragraph.runs:
            paragraph.runs[0].text = replaced
        else:
            paragraph.add_run(replaced)

def _replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in cell.tables:
                _replace_in_table(t, mapping)

def replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.header.tables:
            _replace_in_table(t, mapping)

        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for t in section.footer.tables:
            _replace_in_table(t, mapping)

# ---------- Excel helpers ----------
def normalize(s: str) -> str:
    return str(s).strip().lower()

def find_col(df, candidates):
    cols = {normalize(c): c for c in df.columns}
    for cand in candidates:
        key = normalize(cand)
        if key in cols:
            return cols[key]
    return None

def to_str(x):
    if pd.isna(x):
        return ""
    if hasattr(x, "strftime"):
        return x.strftime("%d/%m/%Y")

    if isinstance(x, float):
        if x.is_integer():
            return str(int(x))
        return str(x)

    if isinstance(x, int):
        return str(x)

    s = str(x).strip()
    if s.endswith(".0") and s.replace(".0", "").isdigit():
        s = s[:-2]
    return s

def safe_filename(s: str) -> str:
    s = (s or "contract").strip()
    for ch in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
        s = s.replace(ch, "-")
    return s

# ---------- DOCX bytes -> PDF bytes (LibreOffice with isolated profile) ----------
def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    if not PDF_AVAILABLE:
        raise RuntimeError("تحويل PDF غير متاح.")

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "file.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        user_profile = os.path.join(tmp, "lo_profile")
        user_install = "file:///" + user_profile.replace("\\", "/")

        cmd = [
            SOFFICE,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation={user_install}",
            "--convert-to", "pdf",
            "--outdir", tmp,
            docx_path,
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed (code {result.returncode}).\n"
                f"STDOUT:\n{result.stdout}\n\nSTDERR:\n{result.stderr}"
            )

        pdf_path = os.path.join(tmp, "file.pdf")
        if not os.path.exists(pdf_path):
            pdf_path = os.path.join(tmp, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")

        with open(pdf_path, "rb") as f:
            return f.read()

# ---------- Uploads ----------
st.subheader("1) ارفع قالب العقد (Word)")
template_file = st.file_uploader("قالب العقد (DOCX)", type=["docx"])

st.subheader("2) ارفع قاعدة البيانات (Excel)")
excel_file = st.file_uploader("قاعدة البيانات (XLSX)", type=["xlsx"])

# ---------- Column configuration ----------
REQUIRED_COLUMNS = {
    "name": ["الاسم", "اسم", "Name", "name"],
    "id": ["رقم الهوية", "هوية رقم", "الهوية", "ID", "id"],
    "nationality": ["الجنسية", "Nationality", "nationality"],
    "address": ["العنوان", "Address", "address"],
    "city": ["المدينة", "المدينه", "City", "city"],
    "email": ["البريد الإلكتروني", "البريد الالكتروني", "Email", "email"],
    "mobile": ["الجوال", "رقم الجوال", "Mobile", "mobile"],
    "birth_date": ["تاريخ الميلاد", "Birth Date", "birth_date"],
}

OPTIONAL_COLUMNS = {
    "birth_place": ["مكان الميلاد", "Birth Place", "birth_place"],
    "expiry_date": ["تاريخ الانتهاء", "تاريخ الإنتهاء", "Expiry Date", "expiry_date"],
    "marital_status": ["الحالة الاجتماعية", "الحاله الاجتماعيه", "Marital Status", "marital_status"],
}

if not (template_file and excel_file):
    st.info("ارفع ملف Word القالب + ملف Excel قاعدة البيانات لتفعيل الاختيار والتوليد.")
    st.stop()

# Read Excel
df = pd.read_excel(excel_file)

# Resolve columns
resolved = {}
missing_required = []
for key, candidates in REQUIRED_COLUMNS.items():
    col = find_col(df, candidates)
    if col is None:
        missing_required.append(key)
    else:
        resolved[key] = col

if missing_required:
    st.error("أعمدة إلزامية ناقصة أو أسماء الأعمدة مختلفة في Excel:")
    st.write(missing_required)
    st.stop()

for key, candidates in OPTIONAL_COLUMNS.items():
    col = find_col(df, candidates)
    if col is not None:
        resolved[key] = col

# Clean rows
df["_name"] = df[resolved["name"]].apply(to_str)
df = df[df["_name"].astype(str).str.strip() != ""].copy()
if df.empty:
    st.error("ملف Excel لا يحتوي أسماء صالحة في عمود الاسم.")
    st.stop()

df["_id"] = df[resolved["id"]].apply(to_str)
df["_display"] = df["_name"] + " | " + df["_id"]

selected = st.selectbox("3) اختر الشخص", df["_display"].tolist())
row = df[df["_display"] == selected].iloc[0]

def get_value(row_obj, key):
    if key in resolved:
        return to_str(row_obj[resolved[key]])
    return ""

def build_mapping(row_obj):
    return {
        "{{name}}": get_value(row_obj, "name"),
        "{{nationality}}": get_value(row_obj, "nationality"),
        "{{id}}": get_value(row_obj, "id"),
        "{{birth_place}}": get_value(row_obj, "birth_place"),
        "{{birth_date}}": get_value(row_obj, "birth_date"),
        "{{expiry_date}}": get_value(row_obj, "expiry_date"),
        "{{marital_status}}": get_value(row_obj, "marital_status"),
        "{{address}}": get_value(row_obj, "address"),
        "{{city}}": get_value(row_obj, "city"),
        "{{email}}": get_value(row_obj, "email"),
        "{{mobile}}": get_value(row_obj, "mobile"),
    }

mapping = build_mapping(row)

st.subheader("معاينة البيانات المختارة")
st.write(mapping)

want_pdf = st.checkbox("توليد PDF أيضًا", value=True)

colA, colB = st.columns(2)

# ---------- Generate single contract ----------
with colA:
    if st.button("إنشاء عقد الشخص المختار", type="primary"):
        doc = Document(io.BytesIO(template_file.getvalue()))
        replace_everywhere(doc, mapping)

        docx_out = io.BytesIO()
        doc.save(docx_out)
        docx_bytes = docx_out.getvalue()

        safe_name = safe_filename(mapping["{{name}}"])
        base_single = f"Contract_{safe_name}"

        st.download_button(
            "تحميل Word",
            data=docx_bytes,
            file_name=base_single + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        if want_pdf:
            try:
                pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)
                st.download_button(
                    "تحميل PDF",
                    data=pdf_bytes,
                    file_name=base_single + ".pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"فشل تحويل PDF: {e}")

# ---------- Generate ALL contracts ZIP ----------
with colB:
    if st.button("إنشاء جميع العقود (ZIP)"):
        zbuf = io.BytesIO()

        with zipfile.ZipFile(zbuf, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for _, r in df.iterrows():
                mp = build_mapping(r)
                if not mp["{{name}}"].strip():
                    continue

                doc = Document(io.BytesIO(template_file.getvalue()))
                replace_everywhere(doc, mp)

                docx_out = io.BytesIO()
                doc.save(docx_out)
                docx_bytes = docx_out.getvalue()

                base = f"Contract_{safe_filename(mp['{{name}}'])}_{safe_filename(mp['{{id}}'])}"
                z.writestr(base + ".docx", docx_bytes)

                if want_pdf:
                    try:
                        pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)
                        z.writestr(base + ".pdf", pdf_bytes)
                    except Exception:
                        pass

        zbuf.seek(0)
        st.download_button(
            "تحميل ZIP (كل العقود)",
            data=zbuf,
            file_name="All_Contracts.zip",
            mime="application/zip"
        )

